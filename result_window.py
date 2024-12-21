import customtkinter as CTk
from tkinter import filedialog, messagebox
import os
from docx import Document
from datetime import datetime


def save_report(check_results, result_window):
    # Отключаем флаг topmost для главного окна (result_window не выходит на передний план)
    result_window.attributes("-topmost", False)

    if check_results is None:
        messagebox.showerror("Ошибка", "Нет результатов для сохранения.", parent=result_window)
        result_window.attributes("-topmost", True)  # Восстанавливаем флаг topmost
        return

    # Получаем текущую дату и время для имени файла
    current_datetime = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
    default_filename = f"Отчёт о проверке {current_datetime}"

    # Открываем диалог выбора файла с дефолтным именем
    file_path = filedialog.asksaveasfilename(
        defaultextension=".txt",
        initialfile=default_filename,  # Устанавливаем дефолтное имя файла
        filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")],
        title="Сохранить отчет"
    )

    if not file_path:  # Если пользователь не выбрал файл
        result_window.attributes("-topmost", True)  # Восстанавливаем флаг topmost
        return

    # Выбираем тип файла по расширению
    file_extension = os.path.splitext(file_path)[1].lower()

    try:
        if file_extension == ".txt":
            # Сохраняем отчет как текстовый файл
            with open(file_path, 'w', encoding='utf-8') as file:
                save_as_txt(check_results, file)
            messagebox.showinfo("Сохранено", "Отчет сохранен как TXT.", parent=result_window)

        elif file_extension == ".docx":
            # Сохраняем отчет как Word документ
            save_as_docx(check_results, file_path)
            messagebox.showinfo("Сохранено", "Отчет сохранен как DOCX.", parent=result_window)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при сохранении файла: {e}", parent=result_window)

    # Восстанавливаем флаг topmost
    result_window.attributes("-topmost", True)


def save_as_txt(check_results, file):
    # Преобразуем результаты проверки в текстовый формат и записываем в файл
    if 'Общие' in check_results:
        file.write("Общие несоответствия:\n")
        for discrepancy in check_results['Общие']:
            file.write(f"- {discrepancy}\n")
        file.write('\n')

    page_keys = [key for key in check_results.keys() if key != 'Общие']
    page_numbers = [key for key in page_keys if isinstance(key, int)]
    sorted_pages = sorted(page_numbers)

    for page in sorted_pages:
        file.write(f"Страница {page}:\n")
        for discrepancy in check_results[page]:
            file.write(f"- {discrepancy}\n")
        file.write('\n')


def save_as_docx(check_results, file_path):
    # Сохраняем результаты в формат DOCX
    doc = Document()

    if 'Общие' in check_results:
        doc.add_heading('Общие несоответствия', level=1)
        for discrepancy in check_results['Общие']:
            doc.add_paragraph(f"- {discrepancy}")
        doc.add_paragraph()  # пустой абзац

    page_keys = [key for key in check_results.keys() if key != 'Общие']
    page_numbers = [key for key in page_keys if isinstance(key, int)]
    sorted_pages = sorted(page_numbers)

    for page in sorted_pages:
        doc.add_heading(f"Страница {page}", level=1)
        for discrepancy in check_results[page]:
            doc.add_paragraph(f"- {discrepancy}")
        doc.add_paragraph()  # пустой абзац

    doc.save(file_path)


def show_result_window(parent, check_results):
    if check_results is None:
        messagebox.showinfo("Нет результатов", "Сначала выполните проверку документа.", parent=parent)
        return

    # Создаем новое окно для отображения результатов
    result_window = CTk.CTkToplevel(parent)
    result_window.title("Результаты проверки")
    result_window.geometry("850x600")
    result_window.configure(fg_color="#E6E6FA")
    result_window.resizable(True, True)

    # Окно будет появляться на переднем плане
    result_window.attributes("-topmost", True)
    result_window.after(100, lambda: result_window.attributes("-topmost", True))

    # Создаем текстовое поле для отображения результатов
    result_textbox = CTk.CTkTextbox(
        master=result_window,
        font=("Arial", 20),  # Размер шрифта 20
        wrap='word',
        fg_color="#E6E6FA",
        border_width=0
    )
    result_textbox.pack(expand=True, fill='both', padx=10, pady=10)

    # Настраиваем тег для заголовков
    tk_text_widget = result_textbox._textbox
    tk_text_widget.tag_configure('heading', font=("Arial", 20, "bold"))

    # Обработка общих несоответствий, если они есть
    if 'Общие' in check_results:
        discrepancies = check_results['Общие']
        result_textbox.insert('end', "Общие несоответствия:\n", 'heading')
        for discrepancy in discrepancies:
            result_textbox.insert('end', f"- {discrepancy}\n")
        result_textbox.insert('end', '\n')

    # Обработка несоответствий по страницам
    # Получаем список ключей-номеров страниц, исключая 'Общие'
    page_keys = [key for key in check_results.keys() if key != 'Общие']
    # Предполагаем, что ключи страниц - это целые числа
    page_numbers = [key for key in page_keys if isinstance(key, int)]
    sorted_pages = sorted(page_numbers)

    for page in sorted_pages:
        discrepancies = check_results[page]  # Используем page как целое число
        result_textbox.insert('end', f"Страница {page}:\n", 'heading')
        for discrepancy in discrepancies:
            result_textbox.insert('end', f"- {discrepancy}\n")
        result_textbox.insert('end', '\n')

    result_textbox.configure(state='disabled')

    # Добавляем кнопку для сохранения отчета
    save_button = CTk.CTkButton(
        master=result_window,
        text="Сохранить отчет",
        width=300,
        height=50,
        fg_color="MediumSlateBlue",
        font=("Arial", 26, "bold"),
        command=lambda: save_report(check_results, result_window)  # Передаем result_window
    )
    save_button.pack(pady=10)
